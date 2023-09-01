import io
import os
import sys

from collections import namedtuple
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from wand.image import Image


class ImgsToWordApp:

    def __init__(self, **kwargs):
        self.init_params()
        self.set_params(**kwargs)

    def init_params(self):
        exc_exts = ["cmd", "com", "cfg", "dll", "doc", "docx", "docm", "exe",\
        "js", "json", "ini", "md", "php", "py", "spec", "torrent", "txt", "tmp",\
        "xls", "xlsx", "xlsm", "xml"]
        imgh_h = 120    
        imgv_h = 160
        info_func=print
        error_func=lambda *args, **kwargs: print(*args, file=sys.stderr, **kwargs)
        split = False
        self.set_params(exc_exts=exc_exts, imgh_h=imgh_h, imgv_h=imgv_h, split=split, info_func=info_func, error_func=error_func)

    def set_params(self, **kwargs):
        for par in ["exc_exts", "imgh_h", "imgv_h", "dir_path", "doc_name", "split", "info_func", "error_func"]:
            if par in kwargs:
                setattr(self, par, kwargs[par])

    def reset(self):
        self.imgsH = list()
        self.imgsV = list()
        self.WordImage = namedtuple("WordImage", "w h data")

    def process(self):
        for fn in os.listdir(self.dir_path):
            if not (os.path.splitext(fn)[1].lstrip(".").lower() in self.exc_exts)\
            and os.path.isfile(fn):
                try:
                    with Image(filename=fn) as img:
                        w = img.width
                        h = img.height
                        self.info_func(f"Processing {fn}: {img.format} ({w} x {h})", end="")
                        img_h = self.imgh_h if (w >= h) else self.imgv_h
                        factor = img_h / h
                        new_w = round(w * factor)
                        self.info_func(f" => {new_w} x {img_h}")
                        img.format = "JPEG"
                        img.resize(new_w, img_h, "catrom", 1)
                        stream = io.BytesIO(img.make_blob())
                        wimg = self.WordImage(new_w, img_h, stream)
                        (self.imgsH if w >= h else self.imgsV).append(wimg)
                except Exception as e:
                    self.error_func(f"ERROR during processing {fn}:", str(e))
                    #raise()

    def process_list(self, lst, document):
        par = document.add_paragraph()
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        run = par.add_run()
        for wimg in lst:
            run.add_picture(wimg.data, width=Pt(wimg.w))
            run = par.add_run(" ")

    def save(self):
        document = Document()
        self.process_list(self.imgsH, document)
        if self.split:
            document.add_page_break()
        self.process_list(self.imgsV, document)
        document.save(self.doc_name)
    
    def convert(self):
        self.reset()
        self.process()
        self.save()
